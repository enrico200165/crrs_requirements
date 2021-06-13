import docx.table

import crrs_reqs_global_defs as g


class MSWordTable:
    """
    Rudimentary wrapper class around import docx.table
    implements only the methods needed for my task
    (getting the leftmost cell content for each row of a word table containing requirements)
    """

    def __init__(self, table):
        if not isinstance(table, docx.table.Table):
            exit(1)
        self._table = table

    def get_1st_cell(self):
        return self._table.rows[0].cells[0].paragraphs[0].text

    def row_left_cell_equals(self, col_name):
        if self._table is None:
            return None
        if self._table.rows is None:
            return None
        for row in self._table.rows:
            leftmost = row.cells[0]
            if leftmost is None:
                return None
            for par in leftmost.paragraphs:
                text = par.text
                if col_name == text:
                    return row
        return None

    def get_cell_row_id(self, row_id, cell_nr):
        row = self.row_left_cell_equals(row_id)
        if row is None:
            return None
        cell = row.cells[cell_nr]
        return cell

    def get_tex_cell_row_id(self, row_id, cell_nr, par_nr):
        pars = self.get_pars_cell_row_id(row_id, cell_nr, par_nr)
        text = "\n".join(pars)
        return text

    def get_pars_cell_row_id(self, row_id, cell_nr, par_nr):
        cell = self.get_cell_row_id(row_id, cell_nr)
        if cell is None:
            return None
        pars = []
        for par in cell.paragraphs:
            pars.append(par.text)
        return pars


