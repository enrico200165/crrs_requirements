"""
Produce a pandas data frame
containing requirement entries and nothing more
hiding the production process (classes, Etc)
"""

import pandas as pd
from docx.api import Document

import crrs_reqs_global_defs as g
import req_entry
from req_entry_word_table_class import ReqEntryWordTable


def ReqEntry_from_WordReqEntry(word_req):
    req = req_entry.ReqEntry(
        word_req.rp_spec_id,
        word_req.req_id,
        word_req.name,
        word_req.description,
        word_req.status,
        word_req.source,
        word_req.rationale,
        word_req.target_subtarget
    )
    return req


def row_left_cell_equals(t, col_name):
    for row in t.rows:
        leftmost = row.cells[0]
        for par in leftmost.paragraphs:
            if col_name == par.text:
                return row
    return None


def is_word_table_req_table(wt):
    if row_left_cell_equals(wt, g.rp_spec_id_val) is None:
        return False
    if row_left_cell_equals(wt, g.req_id_val) is None:
        return False
    if row_left_cell_equals(wt, g.name_val) is None:
        return False

    return True


# ===================================================================


def add_entries_to_df(df, entries):
    for i in range(len(entries)):
        e = entries[i]
        values = e.values_list()
        df.loc[i] = values


def reqs_from_word_file(filepath):

    doc = Document(docx=filepath)

    reqs_l = []
    for i, t in enumerate(doc.tables):
        if not is_word_table_req_table(t):
            continue
        wt = ReqEntryWordTable(t)
        # print(wt.col1_to_string())
        re = ReqEntry_from_WordReqEntry(wt)
        reqs_l.append(re)

    df = pd.DataFrame(columns=g.col_names)
    add_entries_to_df(df, reqs_l)

    # print("read file {}".format(g.filename))
    # print("reading in directory {}".format(g.path))
    return df
